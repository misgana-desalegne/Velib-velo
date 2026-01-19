from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create Document
doc = Document()

# Title
title = doc.add_heading('Projet Vélib - Application Structure', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Subtitle
subtitle = doc.add_paragraph('Comprehensive Bike-Sharing Analytics Platform')
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.italic = True

doc.add_paragraph()

# Project Overview
doc.add_heading('Project Overview', level=1)
doc.add_paragraph(
    'Projet Vélib is a comprehensive bike-sharing analytics platform for Paris\'s Vélib system. '
    'It provides real-time station monitoring, historical data analysis, and actionable insights '
    'through an interactive dashboard.'
)

# Root Directory
doc.add_heading('Root Directory Structure', level=1)
doc.add_paragraph(
    'Projet_velib/',
    style='List Bullet'
)
structures = [
    ('apps/', 'Django applications'),
    ('frontend/', 'React TypeScript frontend'),
    ('projet_velib/', 'Django project configuration'),
    ('build/', 'Compiled frontend build'),
    ('manage.py', 'Django management script'),
    ('package.json', 'Frontend dependencies'),
    ('requirements.txt', 'Backend dependencies'),
    ('vite.config.ts', 'Vite build configuration'),
    ('db.sqlite3', 'SQLite database'),
]
for item, desc in structures:
    doc.add_paragraph(f'{item}: {desc}', style='List Bullet 2')

# Backend Architecture
doc.add_heading('Backend Architecture', level=1)

doc.add_heading('Django Project Configuration (projet_velib/)', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'File'
hdr_cells[1].text = 'Purpose'
rows_data = [
    ('settings.py', 'Django configuration (database, apps, middleware)'),
    ('urls.py', 'Main URL router'),
    ('wsgi.py', 'WSGI server configuration'),
    ('asgi.py', 'ASGI server configuration'),
]
for i, (file, purpose) in enumerate(rows_data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = file
    row_cells[1].text = purpose

doc.add_heading('Analytics App (apps/analytics/)', level=2)

doc.add_heading('Core Files', level=3)
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'File'
hdr_cells[1].text = 'Purpose'
files_data = [
    ('models.py', 'Database models'),
    ('serializers.py', 'API data serializers'),
    ('views.py', 'API view handlers'),
    ('urls.py', 'App URL routing'),
    ('admin.py', 'Django admin configuration'),
    ('apps.py', 'App initialization'),
]
for i, (file, purpose) in enumerate(files_data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = file
    row_cells[1].text = purpose

doc.add_heading('Services (apps/analytics/services/)', level=3)
table = doc.add_table(rows=8, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Module'
hdr_cells[1].text = 'Description'
services_data = [
    ('etl_pipeline.py', 'Main ETL orchestrator'),
    ('extractor.py', 'Vélib API data extraction'),
    ('transformer.py', 'Data cleaning, validation, aggregation'),
    ('loader.py', 'Database loading operations'),
    ('analytics_service.py', 'Analytics calculations'),
    ('advanced_analytics_service.py', 'Advanced metrics & insights'),
    ('arrondissement_service.py', 'District-specific analysis'),
]
for i, (module, desc) in enumerate(services_data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = module
    row_cells[1].text = desc

# Frontend Architecture
doc.add_heading('Frontend Architecture', level=1)

doc.add_heading('Structure (frontend/)', level=2)
frontend_items = [
    ('api/', 'API communication layer'),
    ('features/', 'Feature modules'),
    ('pages/', 'Page components'),
    ('shared/', 'Shared resources'),
    ('assets/', 'Static assets'),
    ('App.tsx', 'Root component'),
    ('main.tsx', 'Entry point'),
]
for item, desc in frontend_items:
    doc.add_paragraph(f'{item}: {desc}', style='List Bullet')

doc.add_heading('Key Technologies', level=2)
tech_items = [
    'Framework: React 18+ with TypeScript',
    'Build Tool: Vite',
    'UI Components: Radix UI',
    'CSS: Tailwind CSS',
    'State Management: React Hooks',
    'API Client: Axios/Fetch',
]
for tech in tech_items:
    doc.add_paragraph(tech, style='List Bullet')

# Database Models
doc.add_heading('Database Models', level=1)
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Model'
hdr_cells[1].text = 'Purpose'
models_data = [
    ('Commune', 'Paris districts/communes'),
    ('BikeStation', 'Bike station locations & info'),
    ('Trip', 'Individual bike trips'),
    ('StationStatus', 'Real-time station snapshots'),
    ('DailyAnalytics', 'Daily aggregated metrics'),
    ('WeeklyAnalytics', 'Weekly aggregated metrics'),
]
for i, (model, purpose) in enumerate(models_data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = model
    row_cells[1].text = purpose

# API Endpoints
doc.add_heading('API Endpoints', level=1)

doc.add_heading('Authentication', level=2)
endpoints = [
    'POST /api/auth/login/ - User login',
    'POST /api/auth/register/ - User registration',
    'POST /api/auth/logout/ - User logout',
    'GET /api/auth/profile/ - Get current user',
]
for endpoint in endpoints:
    doc.add_paragraph(endpoint, style='List Bullet')

doc.add_heading('Analytics', level=2)
analytics_endpoints = [
    'GET /api/analytics/stations/ - List all stations',
    'GET /api/analytics/stations/{id}/ - Station details',
    'GET /api/analytics/communes/ - List communes',
    'GET /api/analytics/daily/ - Daily analytics',
    'GET /api/analytics/weekly/ - Weekly analytics',
    'GET /api/analytics/arrondissement/ - District analysis',
]
for endpoint in analytics_endpoints:
    doc.add_paragraph(endpoint, style='List Bullet')

# Build & Deployment
doc.add_heading('Build & Deployment', level=1)

doc.add_heading('Frontend Build', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Command'
hdr_cells[1].text = 'Purpose'
build_data = [
    ('npm install', 'Install dependencies'),
    ('npm run dev', 'Start development server'),
    ('npm run build', 'Build for production'),
]
for i, (cmd, purpose) in enumerate(build_data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = cmd
    row_cells[1].text = purpose

doc.add_heading('Backend Setup', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Command'
hdr_cells[1].text = 'Purpose'
backend_data = [
    ('pip install -r requirements.txt', 'Install dependencies'),
    ('python manage.py migrate', 'Run database migrations'),
    ('python manage.py runserver', 'Start development server'),
]
for i, (cmd, purpose) in enumerate(backend_data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = cmd
    row_cells[1].text = purpose

# Key Features
doc.add_heading('Key Features', level=1)
features = [
    ('Real-Time Monitoring', [
        'Live bike station status',
        'Available bikes & docks',
        'Station occupancy rates',
    ]),
    ('Historical Analytics', [
        'Trip analysis',
        'Usage patterns',
        'Trend identification',
    ]),
    ('District Analysis', [
        'Arrondissement-level insights',
        'Geographic distribution',
        'Performance metrics',
    ]),
    ('User Management', [
        'Authentication system',
        'Team collaboration',
        'Role-based access',
    ]),
    ('Dashboard', [
        'Interactive visualizations',
        'Real-time updates',
        'Analytics charts',
        'Map view',
    ]),
]

for feature_name, details in features:
    doc.add_heading(feature_name, level=2)
    for detail in details:
        doc.add_paragraph(detail, style='List Bullet')

# Technology Stack
doc.add_heading('Technology Stack Summary', level=1)
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Layer'
hdr_cells[1].text = 'Technology'
tech_stack = [
    ('Frontend', 'React + TypeScript, Vite, Radix UI, Tailwind CSS'),
    ('Backend', 'Django, Django REST Framework, Python'),
    ('Database', 'SQLite'),
    ('Data Processing', 'Pandas, NumPy'),
    ('Scheduling', 'APScheduler (optional)'),
]
for i, (layer, tech) in enumerate(tech_stack, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = layer
    row_cells[1].text = tech

# Development Workflow
doc.add_heading('Development Workflow', level=1)
workflow_steps = [
    'Setup Phase: Install dependencies, run migrations, configure environment',
    'Development Phase: Frontend & backend run in development mode',
    'ETL Phase: ETL pipeline extracts and processes Vélib data',
    'API Phase: Django REST API serves processed data',
    'Frontend Phase: React dashboard consumes and visualizes data',
    'Deployment Phase: Build frontend, containerize, deploy',
]
for i, step in enumerate(workflow_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

# Project Statistics
doc.add_heading('Project Statistics', level=1)
stats = [
    'Backend Routes: ~20+ API endpoints',
    'Database Models: 6+ core models',
    'Frontend Pages: 5+ main pages',
    'Services: 4+ analytical services',
    'Frontend Components: 50+ Radix UI components',
]
for stat in stats:
    doc.add_paragraph(stat, style='List Bullet')

# Save document
doc.save('Projet_Velib_Structure.docx')
print('✓ Word document created successfully: Projet_Velib_Structure.docx')
